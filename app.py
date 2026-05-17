import streamlit as st
import google.generativeai as genai
import os
import io
from docx import Document
from fpdf import FPDF

# --- Google Gemini API Configuration ---
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=GEMINI_API_KEY)

def generate_summary(pdf_path, language):
    """PDF එක කියවා මූලික සාරාංශය සහ එහි ඇති ගණිතමය සමීකරණ ලැයිස්තුව ලබාගැනීම"""
    try:
        uploaded_gemini_file = genai.upload_file(path=pdf_path)
        
        if language == "සිංහල (Sinhala)":
            lang_instruction = "Write the summary clearly in academic Sinhala. Keep mathematical formulas in standard LaTeX notation."
        elif language == "한국어 (Korean)":
            lang_instruction = "Write the summary professionally in academic Korean (합니다/습니다 style). Keep mathematical formulas in standard LaTeX notation."
        else:
            lang_instruction = "Write the summary clearly in academic English."

        prompt = f"""
        You are an expert academic research assistant. Analyze the attached research paper PDF.
        Provide a concise yet informative summary under the following clear sections:
        
        # 📄 Research Paper Summary
        
        ### 🎯 1. Main Objectives & Contributions
        - What is the core problem and proposed solution?
        
        ### ⚙️ 2. Methodology Overview
        - Briefly explain the core system, algorithm, or experimental approach.
        
        ### 🧮 3. Key Mathematical Equations Found
        - List all the critical mathematical formulas, loss functions, or core equations used in this paper.
        - **Format Requirement**: You MUST wrap every equation using standard LaTeX notation ($...$ for inline or $$...$$ for block equations) so they render beautifully.
        
        ### 📊 4. Key Findings & Conclusion
        - What are the major results and takeaways?
        
        ---
        Language Constraint: {lang_instruction}
        """

        model = genai.GenerativeModel('gemini-3-flash-preview')
        response = model.generate_content([uploaded_gemini_file, prompt])
        
        # Clean up file from Gemini server
        genai.delete_file(uploaded_gemini_file.name)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

def explain_math_deeply(pdf_path, language):
    """පරිශීලකයා බොත්තම එබූ විට පමණක් ක්‍රියාත්මක වන, ගණිතමය කොටස් ගැඹුරින් විග්‍රහ කරන ශ්‍රිතය"""
    try:
        uploaded_gemini_file = genai.upload_file(path=pdf_path)
        
        if language == "සිංහල (Sinhala)":
            lang_instruction = "Explain the mathematics step-by-step deeply in clear Sinhala. Break down all variables and symbols."
        elif language == "한국어 (Korean)":
            lang_instruction = "Explain the mathematics step-by-step deeply in clear Korean. Break down all variables and symbols."
        else:
            lang_instruction = "Explain the mathematics step-by-step deeply in English. Break down all variables and symbols."

        prompt = f"""
        You are a professor in mathematics and computer science. Look at the attached research paper PDF.
        Identify all mathematical formulations, derivations, theorems, or formulas mentioned.
        
        Provide a thorough, step-by-step deep dive:
        1. Display each equation beautifully using LaTeX ($...$ or $$...$$).
        2. Break down every single variable, subscript, superscript, and symbol inside the equations.
        3. Explain the intuition behind *why* this specific math/formula is used in this research and how it works logically.
        
        ---
        Language Constraint: {lang_instruction}
        """

        model = genai.GenerativeModel('gemini-3-flash-preview')
        response = model.generate_content([uploaded_gemini_file, prompt])
        
        genai.delete_file(uploaded_gemini_file.name)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

def generate_citation(pdf_path, citation_style):
    """PDF එකෙන් තොරතුරු උකහාගෙන නිවැරදි ආකෘතියට Citation එක සකසන ශ්‍රිතය"""
    try:
        uploaded_gemini_file = genai.upload_file(path=pdf_path)
        
        if citation_style == "IEEE Format":
            style_instruction = "Format the citation strictly following the standard IEEE referencing style guidelines."
        else:
            style_instruction = "Format the citation strictly as a valid BibTeX entry (standard LaTeX format enclosed in a code block)."

        prompt = f"""
        Analyze the attached research paper PDF to extract its publication details.
        Find the Paper Title, Authors, Journal/Conference Name, Year of Publication, Volume, Issue, and Pages if available.
        
        Generate a ready-to-use academic citation based on this rule:
        {style_instruction}
        
        Do not add extra conversational text, just output the citation directly.
        """

        model = genai.GenerativeModel('gemini-3-flash-preview')
        response = model.generate_content([uploaded_gemini_file, prompt])
        
        genai.delete_file(uploaded_gemini_file.name)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# --- Export Utilities (DOCX & PDF හදන ශ්‍රිතයන්) ---
def create_docx(text):
    doc = Document()
    doc.add_heading('Research Paper Summary Output', level=1)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    
    # fpdf2 වලදී unicode නොවන characters බිඳීම වැළැක්වීමට text එක encode/decode කිරීම
    clean_text = text.encode('latin-1', 'ignore').decode('latin-1')
    
    pdf.multi_cell(0, 10, txt=clean_text)
    return pdf.output()

# --- Streamlit Web Interface Configuration ---
st.set_page_config(page_title="Research Summarizer", page_icon="🔬", layout="centered")

# Custom UI Styling (Lassanata UI eka hadanna CSS)
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        font-weight: 600;
    }
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: rgba(250, 250, 250, 0.9);
        color: #555555;
        text-align: center;
        padding: 10px;
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        font-size: 14px;
        border-top: 1px solid #e0e0e0;
        z-index: 100;
    }
    @media (prefers-color-scheme: dark) {
        .footer {
            background-color: rgba(17, 17, 17, 0.9);
            color: #bbbbbb;
            border-top: 1px solid #333333;
        }
    }
    </style>
""", unsafe_allow_html=True)

st.title("🔬 Smart Research Paper Summarizer")
st.write("Upload a research paper to get a clean summary, extract mathematical formulations, and generate citations instantly.")
st.write("---")

# Configuration Controls
st.subheader("Configuration & File Upload")
language_opt = st.selectbox(
    "Select Output Language:",
    ["English", "සිංහල (Sinhala)", "한국어 (Korean)"]
)

uploaded_file = st.file_uploader("Upload Research Paper (PDF):", type="pdf")

if uploaded_file is not None:
    st.success("File uploaded successfully!")
    
    # Session state initialization
    if "summary_text" not in st.session_state:
        st.session_state.summary_text = None
    if "math_text" not in st.session_state:
        st.session_state.math_text = None
    if "citation_text" not in st.session_state:
        st.session_state.citation_text = None
        
    # Temporary disk save
    temp_filename = "temp_summarizer_paper.pdf"
    with open(temp_filename, "wb") as f:
        f.write(uploaded_file.getbuffer())
        
    # Button 1: Generate Core Summary
    if st.button("Generate Summary"):
        with st.spinner("Analyzing document structure... Please wait..."):
            st.session_state.summary_text = generate_summary(temp_filename, language_opt)
            st.session_state.math_text = None 
            st.session_state.citation_text = None

    # Display Core Summary
    if st.session_state.summary_text:
        st.write("---")
        st.markdown(st.session_state.summary_text)
        st.write("---")
        
        # --- Export Document Section (Download Features) ---
        st.subheader("💾 Export Document")
        st.write("Download the generated summary to your local storage.")
        
        export_format = st.radio("Select File Format:", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)
        
        if export_format == "Word (.docx)":
            docx_data = create_docx(st.session_state.summary_text)
            st.download_button(
                label="📥 Download as DOCX",
                data=docx_data,
                file_name="Research_Summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            pdf_data = create_pdf(st.session_state.summary_text)
            st.download_button(
                label="📥 Download as PDF",
                data=pdf_data,
                file_name="Research_Summary.pdf",
                mime="application/pdf"
            )
            
        st.write("---")
        
        # Interactive Option Sections (Two columns for features)
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🧮 Mathematical Breakdown")
            st.write("Get a step-by-step breakdown of formulas.")
            if st.button("Explain Mathematics"):
                with st.spinner("Deconstructing mathematical formulas... Please wait..."):
                    st.session_state.math_text = explain_math_deeply(temp_filename, language_opt)
                    
        with col2:
            st.subheader("📚 Reference & Citation")
            st.write("Generate a quick reference for your bibliography.")
            citation_style = st.radio("Choose Format:", ["IEEE Format", "BibTeX (LaTeX)"])
            if st.button("Generate Citation"):
                with st.spinner("Extracting metadata for citation... Please wait..."):
                    st.session_state.citation_text = generate_citation(temp_filename, citation_style)
                    
        # Display Math Results below if triggered
        if st.session_state.math_text:
            st.write("---")
            st.info("Mathematical Breakdown Output:")
            st.markdown(st.session_state.math_text)
            
        # Display Citation Results below if triggered
        if st.session_state.citation_text:
            st.write("---")
            st.success("Generated Citation:")
            if "BibTeX" in citation_style:
                st.code(st.session_state.citation_text, language="latex")
            else:
                st.markdown(st.session_state.citation_text)

    # File cleanup
    if os.path.exists(temp_filename):
        os.remove(temp_filename)

else:
    st.info("Please upload a PDF file to begin the analysis.")

# --- Professional Footer (Sankalpa Lokuliyanage & KNU Branding) ---
st.markdown(
    """
    <div class="footer">
        <p>Developed by <b>Sankalpa Lokuliyanage</b> | Kyungpook National University</p>
    </div>
    """,
    unsafe_allow_html=True
)